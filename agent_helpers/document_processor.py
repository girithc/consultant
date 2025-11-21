"""
Document Processing Utilities

Handles text extraction from various document formats (PDF, DOCX, TXT, CSV)
and chunks text for embedding and vector storage.
"""

import io
from typing import List, Dict, Any
from dataclasses import dataclass

# Text extraction libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

# Text splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter


@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted text and metadata"""
    text: str
    metadata: Dict[str, Any]
    chunks: List[str] = None


class DocumentProcessor:
    """Handles document text extraction and chunking"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file bytes"""
        if not PdfReader:
            raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file bytes"""
        if not Document:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_bytes: bytes) -> str:
        """Extract text from TXT file bytes"""
        try:
            # Try UTF-8 first, fall back to other encodings
            try:
                return file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                # Try other common encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        return file_bytes.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                raise ValueError("Could not decode text file with common encodings")
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    def extract_text_from_csv(self, file_bytes: bytes) -> str:
        """Extract text from CSV file bytes"""
        import csv
        try:
            csv_text = file_bytes.decode('utf-8')
            csv_file = io.StringIO(csv_text)
            reader = csv.reader(csv_file)
            
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
        except Exception as e:
            raise ValueError(f"Failed to extract text from CSV: {str(e)}")
    
    def process_document(self, file_bytes: bytes, filename: str, metadata: Dict[str, Any] = None) -> ProcessedDocument:
        """
        Process a document by extracting text and creating chunks
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename (used to determine file type)
            metadata: Optional metadata to attach to the document
        
        Returns:
            ProcessedDocument with extracted text and chunks
        """
        if metadata is None:
            metadata = {}
        
        # Determine file type from filename
        filename_lower = filename.lower()
        
        # Extract text based on file type
        if filename_lower.endswith('.pdf'):
            text = self.extract_text_from_pdf(file_bytes)
            metadata['file_type'] = 'pdf'
        elif filename_lower.endswith('.docx'):
            text = self.extract_text_from_docx(file_bytes)
            metadata['file_type'] = 'docx'
        elif filename_lower.endswith('.txt'):
            text = self.extract_text_from_txt(file_bytes)
            metadata['file_type'] = 'txt'
        elif filename_lower.endswith('.csv'):
            text = self.extract_text_from_csv(file_bytes)
            metadata['file_type'] = 'csv'
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Add filename to metadata
        metadata['filename'] = filename
        metadata['text_length'] = len(text)
        
        # Chunk the text
        chunks = self.chunk_text(text)
        
        return ProcessedDocument(
            text=text,
            metadata=metadata,
            chunks=chunks
        )
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Text to chunk
        
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        return self.text_splitter.split_text(text)


# Convenience function for quick processing
def process_file(file_bytes: bytes, filename: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> ProcessedDocument:
    """
    Process a file and extract text with chunking
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        chunk_size: Size of text chunks (default: 1000)
        chunk_overlap: Overlap between chunks (default: 200)
    
    Returns:
        ProcessedDocument with extracted text and chunks
    """
    processor = DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return processor.process_document(file_bytes, filename)
